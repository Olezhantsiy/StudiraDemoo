"""Generate a stage progress report as DOCX from the university template."""
import io
import math
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from django.core.files.base import ContentFile
from lxml import etree

TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "report_templates" / "stage_report_template.docx"
)

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

_DEGREE_LABEL = {
    "MAG": "Магистрант",
    "ASP": "Аспирант",
}

_PUB_STATUS_LABEL = {
    "DRAFT": "Черновик",
    "PENDING": "На рассмотрении",
    "PRINT": "В печати",
    "PUBLISHED": "Опубликована",
    "REJECTED": "Отклонена",
}

# ── Publication category matchers ────────────────────────────────────────────
# Each entry: (label, index_name_keywords, pub_types_or_None)
# A publication matches a section if ANY of its indexes contains a keyword
# OR if its type is in pub_types (when keywords is empty).
# Section order matches template Table 2 rows:
#   row 3  → wos_scopus
#   row 5  → vak
#   row 7  → rinc
#   row 9  → conference  ("Прочие публикации": тезисы/статьи конференций)
#   row 11 → ip          ("Авторские свидетельства, патенты")
#   row 13 → awards      ("Дипломы победителей конкурсов, олимпиад")
_PUB_SECTIONS = [
    # (section_label, index_keywords, type_filter)
    ("wos_scopus",   ["scopus", "wos"],                         None),
    ("vak",          ["вак", "белый список", "white"],          None),
    ("rinc",         ["ринц", "rsci"],                          None),
    ("conference",   [],                                         ["CONFERENCE", "THESIS"]),
    ("ip",           [],                                         None),   # patents – no model type yet
    ("awards",       [],                                         None),   # diplomas – no model type yet
]


def _pub_in_section(pub, index_keywords, type_filter) -> bool:
    """Return True if a publication belongs to this category section."""
    if index_keywords:
        names = {ix.name.lower() for ix in pub.indexes.all()}
        if any(any(kw in n for n in names) for kw in index_keywords):
            return True
    if type_filter and pub.type in type_filter:
        return True
    return False


# ─── text helpers ────────────────────────────────────────────────────────────

def _short_name(user) -> str:
    if not user:
        return ""
    last = user.last_name or ""
    first = (user.first_name or "")[:1]
    middle = (user.middle_name or "")[:1]
    parts = [last]
    if first:
        parts.append(f"{first}.")
    if middle:
        parts.append(f"{middle}.")
    return " ".join(parts).strip()


def _full_name(user) -> str:
    if not user:
        return ""
    return f"{user.last_name or ''} {user.first_name or ''} {user.middle_name or ''}".strip()


def _academic_year(start_date) -> str:
    if not start_date:
        today = date.today()
        y = today.year if today.month >= 9 else today.year - 1
        return f"{y}/{y + 1}"
    y = start_date.year
    return f"{y}/{y + 1}"


def _study_course(start_date) -> str:
    if not start_date:
        return "1"
    today = date.today()
    return str(max(1, math.ceil((today - start_date).days / 365)))


def _latest_sub_text(task) -> str:
    subs = list(task.submissions.order_by("-created_at"))
    return subs[0].text if subs else ""


# ─── paragraph replacement ───────────────────────────────────────────────────

def _replace_in_para(para, mapping: dict) -> None:
    full = "".join(r.text for r in para.runs)
    new_full = full
    for old, new in mapping.items():
        if old in new_full:
            new_full = new_full.replace(old, str(new))
    if new_full != full and para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""


def _replace_in_cell(cell, mapping: dict) -> None:
    for para in cell.paragraphs:
        _replace_in_para(para, mapping)


# ─── XML / table helpers ─────────────────────────────────────────────────────

def _set_tc_text(tc_element, text: str) -> None:
    """Replace text in a table cell XML element."""
    t_elems = tc_element.findall(f".//{{{WNS}}}t")
    if t_elems:
        t_elems[0].text = text
        for t in t_elems[1:]:
            t.text = ""
    else:
        paras = tc_element.findall(f"{{{WNS}}}p")
        if paras:
            r = etree.SubElement(paras[0], f"{{{WNS}}}r")
            t = etree.SubElement(r, f"{{{WNS}}}t")
            t.text = text


def _find_slot_rows(table, placeholder: str, col: int = 0) -> list[int]:
    """Row indices where cells[col] text contains placeholder."""
    result = []
    for i, row in enumerate(table.rows):
        if col < len(row.cells) and placeholder in row.cells[col].text:
            result.append(i)
    return result


def _fill_block_dynamic(
    table,
    slot_indices: list[int],
    col0_texts: list[str],
    col1_texts: list[str],
) -> None:
    """
    Replace fixed template slot rows with a dynamic number of data rows.
    Uses first slot as formatting template; inserts one empty row if no tasks.
    """
    if not slot_indices:
        return

    tbl = table._tbl
    all_trs = tbl.findall(f"{{{WNS}}}tr")
    template_tr = deepcopy(all_trs[slot_indices[0]])
    anchor_tr = all_trs[slot_indices[0] - 1]

    for idx in sorted(slot_indices, reverse=True):
        all_trs[idx].getparent().remove(all_trs[idx])

    pairs = list(zip(col0_texts, col1_texts)) if col0_texts else [("", "")]
    prev_tr = anchor_tr
    for c0, c1 in pairs:
        new_tr = deepcopy(template_tr)
        cells = new_tr.findall(f".//{{{WNS}}}tc")
        if len(cells) >= 2:
            _set_tc_text(cells[0], c0)
            _set_tc_text(cells[1], c1)
        prev_tr.addnext(new_tr)
        prev_tr = new_tr


def _fill_t2_section(table, data_row_idx: int, pubs: list) -> None:
    """
    Replace one placeholder row in Table 2 with dynamic publication rows.
    One publication per row; empty row inserted if section has no publications.
    Processes one section at a time — call in REVERSE row-index order to avoid
    index shifting.

    Table 2 columns: 0=№, 1=title, 2=status, 3=publisher, 4=volume, 5=stage
    """
    tbl = table._tbl
    all_trs = tbl.findall(f"{{{WNS}}}tr")

    template_tr = deepcopy(all_trs[data_row_idx])
    anchor_tr = all_trs[data_row_idx - 1]     # section header row

    # Remove original slot
    all_trs[data_row_idx].getparent().remove(all_trs[data_row_idx])

    items = pubs if pubs else [None]
    prev_tr = anchor_tr
    for pub in items:
        new_tr = deepcopy(template_tr)
        cells = new_tr.findall(f".//{{{WNS}}}tc")
        if len(cells) >= 4:
            if pub:
                publisher = pub.publisher or ""
                status_label = _PUB_STATUS_LABEL.get(pub.status, pub.status)
                _set_tc_text(cells[1], pub.title)
                _set_tc_text(cells[2], status_label)
                _set_tc_text(cells[3], publisher)
            else:
                _set_tc_text(cells[1], "")
                _set_tc_text(cells[2], "")
                _set_tc_text(cells[3], "")
        prev_tr.addnext(new_tr)
        prev_tr = new_tr


# ─── page-break helper ───────────────────────────────────────────────────────

def _add_page_break_before_second_section(doc: Document) -> None:
    """
    Find the heading paragraph that opens the second section (before Table 1),
    remove the filler empty paragraphs preceding it, and add Word's
    'Page break before' property so the section always starts on a new page.

    Strategy: scan between the two tables, track consecutive-empty-paragraph
    blocks; the first non-empty paragraph that follows a block of ≥ 4 empties
    is the section heading.  All empties in that last block are deleted.
    """
    body = doc.element.body
    children = list(body)

    tbl_positions = [i for i, c in enumerate(children) if c.tag.endswith("}tbl")]
    if len(tbl_positions) < 2:
        return

    t0_pos = tbl_positions[0]
    t1_pos = tbl_positions[1]

    current_empty_block: list = []
    heading_elem = None
    empty_elems_to_remove: list = []

    for i in range(t0_pos + 1, t1_pos):
        child = children[i]
        if not child.tag.endswith("}p"):
            current_empty_block = []
            continue

        text = "".join(
            t.text or ""
            for t in child.findall(f".//{{{WNS}}}t")
        ).strip()

        if not text:
            current_empty_block.append(child)
        else:
            if len(current_empty_block) >= 4:
                # All empties in this block are filler — remove them all
                empty_elems_to_remove.extend(current_empty_block)
                heading_elem = child
                break
            current_empty_block = []

    for elem in empty_elems_to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

    if heading_elem is not None:
        pPr = heading_elem.find(f"{{{WNS}}}pPr")
        if pPr is None:
            pPr = etree.Element(f"{{{WNS}}}pPr")
            heading_elem.insert(0, pPr)
        if pPr.find(f"{{{WNS}}}pageBreakBefore") is None:
            pb = etree.Element(f"{{{WNS}}}pageBreakBefore")
            pPr.insert(0, pb)


# ─── main function ───────────────────────────────────────────────────────────

def generate_stage_docx(stage) -> ContentFile:
    """Open the DOCX template, fill all placeholders with stage data, return ContentFile."""
    from research.models import Publication, TaskType

    project = stage.project
    enrollment = project.enrollment
    student = enrollment.student
    supervisor = project.supervisor

    group = getattr(enrollment, "group", None)
    program = group.program if group else None
    department = program.department if program else None
    dept_head = department.head if department else None

    degree_label = _DEGREE_LABEL.get(
        program.degree_level if program else "", "Аспирант"
    )
    program_name = program.full_name if program else ""
    start_date = getattr(enrollment, "start_date", None)
    acad_year = _academic_year(start_date)
    course = _study_course(start_date)
    student_full = _full_name(student)
    student_short = _short_name(student)
    stage_order = str(stage.order)

    # ── Tasks ──────────────────────────────────────────────────────────────
    tasks = list(stage.tasks.prefetch_related("submissions").order_by("id"))
    file_tasks = [t for t in tasks if t.task_type == TaskType.FILE]
    pub_tasks = [t for t in tasks if t.task_type == TaskType.PUBLICATION]

    # ── Publications ───────────────────────────────────────────────────────
    all_pubs = list(
        Publication.objects.filter(task__stage=stage)
        .prefetch_related("indexes")
        .order_by("id")
    )

    # ── Open template ──────────────────────────────────────────────────────
    doc = Document(str(TEMPLATE_PATH))

    # ── Paragraph placeholders ─────────────────────────────────────────────
    _acad_year_placeholder = (
        "{studentEnrollment.start_date}"
        "/{studentEnrollment.start_date+relativedelta(years=1)}"
    )
    para_mapping = {
        "{studentEnrollment.user.get_full_name}": student_full,
        "{{studentEntolment.user.get_full_name}}": student_full,
        "{ФИО studentEnrollment.user.fullname}": student_full,
        "{studentEnrollment.user.shortname}": student_short,
        "{{studentEnrollment.degreelevel}}": degree_label,
        "{{program}}": program_name,
        "{{ReasearchProject.name}}": project.title,
        "{ReasarchProject,superviser.user.shortname}": _short_name(supervisor),
        "{ReasarchProject.StudentEnrollment.program.departament.head.user.shortname}": _short_name(dept_head),
        # Specific variant with trailing "/учебный год" (para 1 of first page)
        f"{_acad_year_placeholder}/учебный год": f"{acad_year} учебный год",
        # Generic variant — covers "уч.г:", "уч.г.", " учебный год" (space) etc.
        # Must come AFTER the specific key so it doesn't consume it first.
        _acad_year_placeholder: acad_year,
        "{{studentEnrollment.course}}": course,
        "{{ReaseachStage.order}}": stage_order,
    }
    for para in doc.paragraphs:
        _replace_in_para(para, para_mapping)

    # ── Table 1: tasks — dynamic rows ──────────────────────────────────────
    if doc.tables:
        t1 = doc.tables[0]

        def _task_col0(task) -> str:
            return task.description.strip() if task.description.strip() else task.title

        # FILE task block (search col 0)
        file_slots = _find_slot_rows(t1, "{{StageTask.description}}", col=0)
        _fill_block_dynamic(
            t1, file_slots,
            col0_texts=[_task_col0(t) for t in file_tasks],
            col1_texts=[_latest_sub_text(t) for t in file_tasks],
        )

        # PUBLICATION task block (re-scan after FILE modification)
        pub_slots = _find_slot_rows(t1, "{{StageTask(taskType=PUBLICATION )}}", col=0)
        _fill_block_dynamic(
            t1, pub_slots,
            col0_texts=[_task_col0(t) for t in pub_tasks],
            col1_texts=[_latest_sub_text(t) for t in pub_tasks],
        )

    # ── Table 2: publications — per-section dynamic rows ───────────────────
    # A publication can appear in MULTIPLE sections if its indexes match
    # several categories (e.g. an article indexed in both Scopus and VAK).
    # Data rows sit immediately after each section-header row:
    #   row 2 (WoS/Scopus)    → row 3 data slot
    #   row 4 (VAK)           → row 5 data slot
    #   row 6 (RINC)          → row 7 data slot
    #   row 8 (Прочие публик.)→ row 9 data slot   ← conference/THESIS
    #   row 10 (Патенты)      → row 11 data slot  ← ip
    #   row 12 (Дипломы)      → row 13 data slot  ← awards
    # Process in REVERSE order so index shifting doesn't affect earlier sections.
    if len(doc.tables) > 1:
        t2 = doc.tables[1]

        # Original data row indices for the 6 sections (reverse order)
        SECTION_DATA_ROWS = [13, 11, 9, 7, 5, 3]
        sections_reversed = list(reversed(_PUB_SECTIONS))   # awards, ip, conference, rinc, vak, wos

        for data_row_idx, (_, index_kw, type_filter) in zip(
            SECTION_DATA_ROWS, sections_reversed
        ):
            section_pubs = [
                p for p in all_pubs
                if _pub_in_section(p, index_kw, type_filter)
            ]
            _fill_t2_section(t2, data_row_idx, section_pubs)

    # ── Page break before the second section heading ───────────────────────
    _add_page_break_before_second_section(doc)

    # ── Save ───────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"stage_report_{project.id}_{stage.id}.docx"
    return ContentFile(buf.read(), name=filename)
